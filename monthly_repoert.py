import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- GLOBAL CONFIG ---
from monthly_final_report import mothly_final_report

FONT_NAME = 'Nirmala UI'  # Highly recommended for Marathi/English balance

def monthly_repo():
    # --- UTILS ---
    def create_letter_header(doc, phc, upkendra, tal, dist, date, subject, g_date="", p_date=""):
        """Creates the standard PHC letterhead with balanced font rendering."""

        # 1. PHC Details - Top Right Corner
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_text = (
            f"प्रा.आ. केंद्र :- {phc}\n"
            f"उपकेंद्र :- {upkendra}\n"
            f"ता. {tal} जि. {dist}\n"
            f"जा.क्रं ____________\n"
            f"दिनांक:- {date}"
        )
        run = header_para.add_run(header_text)
        run.font.name = FONT_NAME
        run.font.size = Pt(13)
        run.bold = True

        # 2. Recipient Section - Left
        p2 = doc.add_paragraph("\nप्रति\nवरिष्ठ भूवैज्ञानिक\nभूजल सर्वेक्षण आणि विकास यंत्रणा\nउपविभागीय प्रयोगशाळा, इंदापूर")
        for run in p2.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
            run.bold = True

        # 3. Subject
        subj = doc.add_paragraph(f"\nविषय : {subject}")
        subj.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_subj = subj.runs[0]
        run_subj.font.name = FONT_NAME
        run_subj.font.size = Pt(13)
        run_subj.bold = True

        # 4. Body Text
        body = doc.add_paragraph(
            f"उपरोक्त विषयानुसार प्रा.आ. केंद्र {phc}, उपकेंद्र {upkendra} कक्षेतील नमुने तपासणीसाठी पाठवीत आहोत. तरी कृपया तपासून अहवाल मिळावा ही विनंती.\n"
        )
        for run in body.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(12)

        # 5. Dates
        if g_date or p_date:
            date_para = doc.add_paragraph()
            date_line = f"नमुने घेतल्याचा दिनांक: {g_date}" + (" " * 12) + f"नमुने पाठवल्याचा दिनांक: {p_date}"
            run_date = date_para.add_run(date_line)
            run_date.font.name = FONT_NAME
            run_date.font.size = Pt(12)
            run_date.bold = True

    # --- APP LAYOUT ---
    st.set_page_config(page_title="PHC Report Generator", layout="wide")

    tab_water, tab_tcl, tab_salt, tab_monthfinal = st.tabs([
        "💧 पाणी नमुने", "🧪 TCL नमुने", "🧂 मीठ नमुने", "Monthly Final Report"
    ])

    with tab_monthfinal:
        mothly_final_report()

    # --- TAB 1: पाणी नमुने ---
    with tab_water:
        st.header("अणुजैविक/रासायनिक पाणी नमुने तपासणी")
        col1, col2, col3 = st.columns(3)
        with col1:
            w_phc = st.text_input("प्रा.आ. केंद्र", value="शेळगाव", key="w1")
            w_up = st.text_input("उपकेंद्र", value="शेळगाव", key="w_up")
        with col2:
            w_tal = st.text_input("तालुका", value="इंदापूर", key="w2")
            w_dist = st.text_input("जिल्हा", value="पुणे", key="w3")
        with col3:
            w_date = st.text_input("आजचा दिनांक", value="19/1/2026", key="w4")
            w_g_date = st.text_input("पाणी नमुने घेतल्याचा दिनांक", value="19/1/2026", key="wg")
            w_p_date = st.text_input("पाणी नमुने पाठवल्याचा दिनांक", value="19/1/2026", key="wp")

        if 'water_rows' not in st.session_state:
            st.session_state.water_rows = [{'uid': '', 'gp': '', 'wadi': '', 'strot': ''}]

        def add_water_row():
            st.session_state.water_rows.append({'uid': '', 'gp': '', 'wadi': '', 'strot': ''})

        for i, row in enumerate(st.session_state.water_rows):
            c1, c2, c3, c4 = st.columns(4)
            st.session_state.water_rows[i]['uid'] = c1.text_input(f"UID {i + 1}", value=row['uid'], key=f"w_uid_{i}")
            st.session_state.water_rows[i]['gp'] = c2.text_input(f"ग्रामपंचायत {i + 1}", value=row['gp'], key=f"w_gp_{i}")
            st.session_state.water_rows[i]['wadi'] = c3.text_input(f"वाडी/वस्ती {i + 1}", value=row['wadi'], key=f"w_wadi_{i}")
            st.session_state.water_rows[i]['strot'] = c4.text_input(f"स्त्रोत {i + 1}", value=row['strot'], key=f"w_strot_{i}")

        st.button("Add Row", on_click=add_water_row, key="btn_w")

        if st.button("Generate Water Letter"):
            doc = Document()
            create_letter_header(doc, w_phc, w_up, w_tal, w_dist, w_date, "अणुजैविक/रासायनिक पाणी नमुने तपासणी बाबत...", w_g_date, w_p_date)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            for j, text in enumerate(["अ.क्र", "UID", "ग्रामपंचायत", "वाडी/वस्ती", "स्त्रोत"]):
                run = table.rows[0].cells[j].paragraphs[0].add_run(text)
                run.font.name = FONT_NAME
                run.font.size = Pt(12)
                run.bold = True

            for idx, r in enumerate(st.session_state.water_rows):
                row_cells = table.add_row().cells
                row_data = [str(idx + 1), r['uid'], r['gp'], r['wadi'], r['strot']]
                for j, val in enumerate(row_data):
                    run = row_cells[j].paragraphs[0].add_run(val)
                    run.font.name = FONT_NAME
                    run.font.size = Pt(12)

            bio = io.BytesIO()
            doc.save(bio)
            st.download_button("Download Water Letter", data=bio.getvalue(), file_name="Pani_Namune.docx")

    # --- TAB 2: TCL नमुने ---
    with tab_tcl:
        st.header("TCL नमुने तपासणी")
        col1, col2, col3 = st.columns(3)
        with col1:
            t_phc = st.text_input("प्रा.आ. केंद्र", value="शेळगाव", key="t1")
            t_up = st.text_input("उपकेंद्र", value="शेळगाव", key="t_up")
        with col2:
            t_date = st.text_input("आजचा दिनांक", value="19/1/2026", key="t2")
            t_g_date = st.text_input("TCL नमुने घेतल्याचा दिनांक", value="19/1/2026", key="tg")
        with col3:
            t_p_date = st.text_input("TCL नमुने पाठवल्याचा दिनांक", value="19/1/2026", key="tp")

        if 'tcl_rows' not in st.session_state:
            st.session_state.tcl_rows = [{'gp': '', 'company': '', 'batch': '', 'mfg': ''}]

        def add_tcl_row():
            st.session_state.tcl_rows.append({'gp': '', 'company': '', 'batch': '', 'mfg': ''})

        for i, row in enumerate(st.session_state.tcl_rows):
            c1, c2, c3, c4 = st.columns(4)
            st.session_state.tcl_rows[i]['gp'] = c1.text_input(f"ग्रामपंचायत {i + 1}", key=f"t_gp_{i}")
            st.session_state.tcl_rows[i]['company'] = c2.text_input(f"कंपनीचे नाव {i + 1}", key=f"t_co_{i}")
            st.session_state.tcl_rows[i]['batch'] = c3.text_input(f"बॅच नंबर {i + 1}", key=f"t_ba_{i}")
            st.session_state.tcl_rows[i]['mfg'] = c4.text_input(f"MFG Date {i + 1}", key=f"t_mf_{i}")

        st.button("Add Row", on_click=add_tcl_row, key="btn_t")

        if st.button("Generate TCL Letter"):
            doc = Document()
            create_letter_header(doc, t_phc, t_up, "इंदापूर", "पुणे", t_date, "TCL नमुने तपासणी बाबत...", t_g_date, t_p_date)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            headers = ["अ.क्रं", "ग्रामपंचायतचे नाव", "कंपनीचे नाव", "बॅच नंबर", "MFG Date"]
            for j, h in enumerate(headers):
                run = table.rows[0].cells[j].paragraphs[0].add_run(h)
                run.font.name = FONT_NAME
                run.font.size = Pt(12)
                run.bold = True
            for idx, r in enumerate(st.session_state.tcl_rows):
                row_cells = table.add_row().cells
                row_data = [str(idx + 1), r['gp'], r['company'], r['batch'], r['mfg']]
                for j, val in enumerate(row_data):
                    run = row_cells[j].paragraphs[0].add_run(val)
                    run.font.name = FONT_NAME
                    run.font.size = Pt(12)
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button("Download TCL Letter", data=bio.getvalue(), file_name="TCL_Namune.docx")

    # --- TAB 3: मीठ नमुने ---
    with tab_salt:
        st.header("मीठ नमुने तपासणी")
        col1, col2, col3 = st.columns(3)
        with col1:
            s_phc = st.text_input("प्रा.आ. केंद्र", value="शेळगाव", key="s1")
            s_up = st.text_input("उपकेंद्र", value="शेळगाव", key="s_up")
        with col2:
            s_date = st.text_input("आजचा दिनांक", value="19/1/2026", key="s2")
            s_g_date = st.text_input("मीठ नमुने घेतल्याचा दिनांक", value="19/1/2026", key="sg")
        with col3:
            s_p_date = st.text_input("मीठ नमुने पाठवल्याचा दिनांक", value="19/1/2026", key="sp")

        if 'salt_rows' not in st.session_state:
            st.session_state.salt_rows = [{'shop': '', 'village': '', 'company': '', 'batch': '', 'mfg': ''}]

        def add_salt_row():
            st.session_state.salt_rows.append({'shop': '', 'village': '', 'company': '', 'batch': '', 'mfg': ''})

        for i, row in enumerate(st.session_state.salt_rows):
            c1, c2, c3, c4, c5 = st.columns(5)
            st.session_state.salt_rows[i]['shop'] = c1.text_input(f"किराणा दुकान {i + 1}", key=f"s_sh_{i}")
            st.session_state.salt_rows[i]['village'] = c2.text_input(f"गावाचे नाव {i + 1}", key=f"s_vi_{i}")
            st.session_state.salt_rows[i]['company'] = c3.text_input(f"कंपनीचे नाव {i + 1}", key=f"s_co_{i}")
            st.session_state.salt_rows[i]['batch'] = c4.text_input(f"बॅच नंबर {i + 1}", key=f"s_ba_{i}")
            st.session_state.salt_rows[i]['mfg'] = c5.text_input(f"MFG Date {i + 1}", key=f"s_mf_{i}")

        st.button("Add Row", on_click=add_salt_row, key="btn_s")

        if st.button("Generate Salt Letter"):
            doc = Document()
            create_letter_header(doc, s_phc, s_up, "इंदापूर", "पुणे", s_date, "मीठ नमुने तपासणी बाबत...", s_g_date, s_p_date)
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            headers = ["अ.क्रं", "किराणा दुकान", "गावाचे नाव", "कंपनीचे नाव", "Batch No", "MFG Date"]
            for j, h in enumerate(headers):
                run = table.rows[0].cells[j].paragraphs[0].add_run(h)
                run.font.name = FONT_NAME
                run.font.size = Pt(12)
                run.bold = True
            for idx, r in enumerate(st.session_state.salt_rows):
                row_cells = table.add_row().cells
                row_data = [str(idx + 1), r['shop'], r['village'], r['company'], r['batch'], r['mfg']]
                for j, val in enumerate(row_data):
                    run = row_cells[j].paragraphs[0].add_run(val)
                    run.font.name = FONT_NAME
                    run.font.size = Pt(12)
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button("Download Salt Letter", data=bio.getvalue(), file_name="Mith_Namune.docx")


if __name__ == "__main__":
    monthly_repo()